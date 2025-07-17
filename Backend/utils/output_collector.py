# ==== IMPORTS ====
import re
import json
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import sys
import os

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from utils.rag_standards import search_standard
from tools.Task_3n11.thuyet_minh_agent import rag_for_task12
# ==== TASK TABLE POSITIONS ====
TASK_CELL_POSITIONS = {
    "task2": {"row": 8, "col_noi_dung": 2, "col_ket_luan": 5},
    "task3": {"row": 10, "col_noi_dung": 2, "col_ket_luan": 5},
    "task5": {"row": 12, "col_noi_dung": 2, "col_ket_luan": 5},
    "task6": {"row": 13, "col_noi_dung": 2, "col_ket_luan": 5},
    "task11": {"row": 20, "col_noi_dung": 2, "col_ket_luan": 5},
    "task4": {"row": 11, "col_noi_dung": 2, "col_ket_luan": 5},
    "task13": {"row": 23, "col_noi_dung": 2, "col_ket_luan": 5},
    "task1": {"row": 5, "col_noi_dung": 2, "col_ket_luan": 5},
    "task12": {"row": 21, "col_noi_dung": 2, "col_ket_luan": 5},
    "task9": {"row": 16, "col_noi_dung": 2, "col_ket_luan": 5}
    
}

# ==== UTILS ====
def extract_number(text, keyword):
    pattern = rf"{keyword}[^\d]*(\d+[\.,]?\d*)"
    match = re.search(pattern, text, re.IGNORECASE)
    return match.group(1) if match else None

# ==== MAIN FUNCTION ====
def output_collector(result_text, word_template_path, word_output_path, llm_compare_func=None, task=None):
    doc = Document(word_template_path)
    table = max(doc.tables, key=lambda t: len(t.rows)*len(t.columns))
    pos = TASK_CELL_POSITIONS.get(task, {"row": 10, "col_noi_dung": 2, "col_ket_luan": 5})
    row = pos["row"]
    col_noi_dung_thiet_ke = pos["col_noi_dung"]  
    col_ket_luan = pos["col_ket_luan"]
    if len(table.rows) <= row or len(table.columns) <= max(col_noi_dung_thiet_ke, col_ket_luan):
        print(f"[output_collector] Bảng không đủ hàng/cột: rows={len(table.rows)}, cols={len(table.columns)}")
        return
    # Ghi kết quả tool vào cột 'Nội dung thiết kế'
    cell = table.cell(row, col_noi_dung_thiet_ke)
    cell.text = result_text
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
            run.font.size = Pt(11)

    # Lấy quy chuẩn từ RAG
    #rag_standard = search_standard(task=task)
    #quy_chuan_text = rag_standard["standard"] if rag_standard else ""

    # Nếu có hàm LLM đối chiếu thì sinh kết luận và ghi vào cột 'Kết luận'
    if llm_compare_func is not None:
        try:
            ket_luan = llm_compare_func(result_text), #quy_chuan_text )
            print(f"[DEBUG] Kết luận sinh bởi LLM: {ket_luan}")
        except Exception as e:
            ket_luan = f"Lỗi gọi LLM: {e}"
            print(f"[DEBUG] Lỗi khi gọi LLM: {e}")
        try:
            cell2 = table.cell(row, col_ket_luan)
            cell2.text = ket_luan
            print(f"[DEBUG] Đã ghi kết luận vào bảng tại dòng {row + 1}, cột {col_ket_luan + 1}")
            # Đặt font Times New Roman, cỡ 11 cho cột Kết luận
            for p in cell2.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
                    run.font.size = Pt(11)
        except Exception as e:
            print(f"[DEBUG] Lỗi khi ghi kết luận vào bảng: {e}")

    try:
        doc.save(word_output_path)
        print(f"[DEBUG] Đã lưu file kết quả: {word_output_path}")
    except Exception as e:
        print(f"[DEBUG] Lỗi khi lưu file kết quả: {e}")

def llm_compare_func_gemini(result_text, api_key, model="models/gemini-1.5-pro"): #standard_text,
    """
    So sánh kết quả với quy chuẩn bằng Gemini, trả về kết luận.
    """
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ KẾT LUẬN:
    - Nếu về tải trọng nền đường, nếu nội dung thiết kế có ghi số liệu về tải trọng nền nền đường cho xe, bãi đỗ nhưng không có thông tin về yêu cầu của cơ quan cảnh sát phòng cháy chữa cháy thì kết luận nguyên văn là "Có ghi số liệu về tải trọng nền đường cho xe, bãi đỗ (ghi số liệu cụ thể như nội dung thiết kế nếu có) nhưng chưa rõ yêu cầu của cơ quan cảnh sát phòng cháy chữa cháy nên chưa thể kết luận." Nếu không có thông tin về tải trọng nền đường thì kết luận nguyên văn là "Không đạt vì không có thông tin về tải trọng nền đường cho xe, bãi đỗ." Nếu có thông tin về tải trọng nền đường cho xe, bãi đỗ và có yêu cầu của cơ quan cảnh sát phòng cháy chữa cháy thì so sánh rồi tự ghi kết luận ngắn gọn. 
    - Nếu về khoảng cách giữa từ mép đường tới tường nhà hoặc công trình, nếu nội dung thiết kế có ghi khoảng cách này thì kết luận nguyên văn là "Khoảng cách từ mép đường tới tường nhà hoặc công trình là (ghi số liệu cụ thể như trong nội dung thiết kế nếu có)." Nếu không có thông tin về khoảng cách này thì kết luận nguyên văn là "Không đạt vì không có thông tin về khoảng cách từ mép đường tới tường nhà hoặc công trình."
    - Nếu về lưu lượng nước, nếu nội dung thiết kế có ghi số liệu về lưu lượng nước chữa cháy và các loại lưu lượng nước khác (nước dùng sinh hoạt...) thì kết luận nguyên văn là "Đạt vì lưu lượng nước chữa cháy đã được đảm bảo ngay cả khi xét đến lưu lượng dùng nước khác."
    - Nếu về độ dốc của đường, nếu nội dung thiết kế có ghi số liệu về độ dốc của các đường và không được vượt quá 1:8,3 thì kết luận nguyên văn là "Đạt vì độ dốc của các đường (số liệu liệt kê như trong nội dung thiết kế nếu có) không vượt quá 1:8,3." Nếu độ dốc vượt quá 1:8,3 thì kết luận nguyên văn là "Không đạt vì độ dốc của các đường (số liệu liệt kê như trong nội dung thiết kế nếu có) vượt quá 1:8,3."
    - Nếu về hệ thống thông tin liên lạc, nếu nội dung thiết kế thể hiện là có bố trí đường dây, có bố trí tủ điện thì kết luận nguyên văn là "Đạt vì có bố trí đường dây, tủ điện." Nếu không có bố trí đường dây, tủ điện thì kết luận nguyên văn là "Không đạt vì không có bố trí đường dây, tủ điện."
    - Nếu về đoạn tránh xe, nếu nội dung thiết kế có đoạn tránh xe hợp lệ thì kết luận nguyên văn là "Đạt vì các đoạn tránh xe hợp lệ." Nếu có các đoạn đường nội bộ cần kiểm tra hoặc có đoạn tránh xe không hợp lệ thì kết luận nguyên văn là "Không đạt vì có đoạn đường nội bộ cần có đoạn tránh xe nhưng chưa được bố trí hoặc có đoạn tránh xe không hợp lệ."
    - Nếu về trụ cứu hỏa, nếu nội dung thiết kế thể hiện có trụ cứu hỏa và khoảng cách giữa các trụ cứu hỏa thì kết luận nguyên văn là "Đạt vì có trụ cứu hỏa và khoảng cách giữa các trụ cứu hỏa là (ghi số liệu cụ thể như trong nội dung thiết kế)." Còn nếu không đủ trụ cứu hỏa để tính khoảng cách thì kết luận nguyên văn là "Có bố trí một trụ cứu hỏa."
    
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"

def llm_compare_func_gemini_task1(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}".
    
    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:
    Nếu nội dung thiết kế không có thông tin về bậc chịu lửa và cấp nguy hiểm cháy của nhà thì kết luận nguyên văn là "Chưa thể kết luận đạt hay không đạt vì chưa rõ thông tin về bậc chịu lửa và cấp nguy hiểm cháy của nhà."
   """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"
    
def llm_compare_func_gemini_task2(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}".

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:
    Nếu nội dung thiết kế không có thông tin về chiều cao phòng cháy chữa cháy thì kết luận nguyên văn là "Chưa thể kết luận bãi đỗ xe chữa cháy đạt hay không đạt vì thiếu thông tin về chiều cao phòng cháy chữa cháy của nhà."
   """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"
    
def llm_compare_func_gemini_task3(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:

    Nếu nội dung thiết kế có ghi số liệu về tải trọng nền nền đường cho xe, bãi đỗ nhưng không có thông tin về yêu cầu của cơ quan cảnh sát phòng cháy chữa cháy thì kết luận nguyên văn là "Có ghi số liệu về tải trọng nền đường cho xe, bãi đỗ (ghi số liệu cụ thể như nội dung thiết kế nếu có) nhưng chưa rõ yêu cầu của cơ quan cảnh sát phòng cháy chữa cháy nên chưa thể kết luận." Nếu không có thông tin về tải trọng nền đường thì kết luận nguyên văn là "Không đạt vì không có thông tin về tải trọng nền đường cho xe, bãi đỗ." Nếu có thông tin về tải trọng nền đường cho xe, bãi đỗ và có yêu cầu của cơ quan cảnh sát phòng cháy chữa cháy thì so sánh rồi tự ghi kết luận ngắn gọn. 
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"

def llm_compare_func_gemini_task4(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}".
    Cách kiểm tra khoảng cách từ mép đường tới tường nhà hoặc công trình có đạt hay không là: Nếu khoảng cách đó không lớn hơn 10m thì là đạt, nếu lớn hơn 10m thì không đạt.

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN: 

    Nếu nội dung thiết kế có ghi khoảng cách từ mép đường tới tường nhà hoặc công trình và khoảng cách này đạt thì kết luận nguyên văn là "Khoảng cách từ mép đường tới tường nhà hoặc công trình là (ghi số liệu cụ thể như trong nội dung thiết kế nếu có) và đạt yêu cầu." 
    Nếu nội dung thiết kế có ghi khoảng cách từ mép đường tới tường nhà hoặc công trình và khoảng cách này không đạt thì kết luận nguyên văn là "Khoảng cách từ mép đường tới tường nhà hoặc công trình là (ghi số liệu cụ thể như trong nội dung thiết kế nếu có) và không đạt yêu cầu."
    Nếu không có thông tin về khoảng cách này thì kết luận nguyên văn là "Không đạt vì không có thông tin về khoảng cách từ mép đường tới tường nhà hoặc công trình."
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"
    
def llm_compare_func_gemini_task5(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:

    Nếu nội dung thiết kế có ghi số liệu về độ dốc của các đường và không được vượt quá 1:8,3 thì kết luận nguyên văn là "Đạt vì độ dốc của các đường (số liệu liệt kê như trong nội dung thiết kế nếu có) không vượt quá 1:8,3." Nếu độ dốc vượt quá 1:8,3 thì kết luận nguyên văn là "Không đạt vì độ dốc của các đường (số liệu liệt kê như trong nội dung thiết kế nếu có) vượt quá 1:8,3."
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"
    
def llm_compare_func_gemini_task6(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:

    Nếu nội dung thiết kế thể hiện kiểu như: "Không có đoạn tránh xe nào được phát hiện." thì kết luận nguyên văn là "Chưa thể kết luận."
    Nếu nội dung thiết kế thể hiện kiểu như: "Đoạn tránh xe hợp lệ ... | Chiều dài: ... | Chiều rộng: ..." hoặc "Không có đoạn tránh xe không hợp lệ nào." thì kết luận nguyên văn là "Đạt."
    Nếu nội dung thiết kế thể hiện kiểu như: "Đoạn tránh xe không hợp lệ ... | Chiều dài: ... | Chiều rộng: ..." hoặc "Không có đạon tránh xe hợp lệ nào." thì kết luận nguyên văn là "Không đạt."

    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"

def llm_compare_func_gemini_task9(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:

    Nếu nội dung thiết kế thể hiện có trụ cứu hỏa thì kết luận nguyên văn là "Đạt vì có trụ cứu hỏa và khoảng cách giữa các trụ cứu hỏa là (ghi số liệu cụ thể như trong nội dung thiết kế)." 
    Nếu không đủ trụ cứu hỏa để tính khoảng cách thì kết luận nguyên văn là "Có bố trí một trụ cứu hỏa."
    Còn nếu không có trụ cứu hỏa nào thì kết luận nguyên văn là "Không đạt vì không có trụ cứu hỏa."
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"

def llm_compare_func_gemini_task11(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:

    Nếu nội dung thiết kế có ghi số liệu về lưu lượng nước chữa cháy và các loại lưu lượng nước khác (nước dùng sinh hoạt...) thì kết luận nguyên văn là "Đạt vì lưu lượng nước chữa cháy đã được đảm bảo ngay cả khi xét đến lưu lượng dùng nước khác."
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"

def llm_compare_func_gemini_task12(result_text, api_key, model="models/gemini-1.5-pro"): 

    prompt = f"""
    Bạn hãy dựa vào diện tích của tòa nhà trong nội dung thiết kế để tính số lượng đám cháy rồi ghi kết luận.
    Nội dung thiết kế: "{result_text}".
    Dữ liệu về dân số: "{rag_for_task12}".
    Tính số lượng đám cháy tính toán theo quy chuẩn sau: Nếu diện tích dưới 1500000 m2 thì tính là 1 đám cháy. Sau đó xét tiếp đến dân số, nếu dân số dưới 10000 người thì tính 1 đám cháy, còn nếu dân số từ 10000 đến 25000 người thì tính là 2 đám cháy. Cuối cùng tổng số đám cháy sẽ là tổng của số đám cháy tính theo diện tích và số đám cháy tính theo dân số. Nếu không có dữ liệu về dân số thì kết luận nguyên văn là "Không có dữ liệu về dân số nên không thể tính số lượng đám cháy theo dân số." Nếu không có diện tích thì kết luận nguyên văn là "Không có diện tích nên không thể tính số lượng đám cháy theo diện tích."
    Sau đó kết luận nguyên văn là: "Số lượng đám cháy tính toán là: ..."
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"

def llm_compare_func_gemini_task13(result_text, api_key, model="models/gemini-1.5-pro"): 
    prompt = f"""
    Bạn là chuyên gia kiểm định xây dựng. Hãy kiểm tra nội dung thiết kế sau và PHẢI đưa ra kết luận theo định dạng trong dấu "..." như ví dụ ở dưới, (nếu có số liệu thì ghi vào) không được ghi thêm:

    Nội dung thiết kế: "{result_text}"

    ĐÂY LÀ VÍ DỤ VỀ CÁCH VIẾT KẾT LUẬN:

    Nếu nội dung thiết kế thể hiện là có bố trí đường dây, có bố trí tủ điện thì kết luận nguyên văn là "Đạt vì có bố trí đường dây, tủ điện." Nếu không có bố trí đường dây, tủ điện thì kết luận nguyên văn là "Không đạt vì không có bố trí đường dây, tủ điện."
    """
    url = f"https://generativelanguage.googleapis.com/v1beta/{model}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    data = {"contents": [{"parts": [{"text": prompt}]}]}
    resp = requests.post(url, headers=headers, data=json.dumps(data))
    if resp.status_code == 200:
        return resp.json()["candidates"][0]["content"]["parts"][0]["text"].strip()
    else:
        return f"Lỗi gọi Gemini: {resp.status_code} - {resp.text}"